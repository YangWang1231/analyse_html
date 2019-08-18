#!/usr/bin/python
 #coding:utf-8
"""
 analyse_html_matrix.py主要作用是获取testbed分析出的质量度量结果，作为生成word报告的数据素材
 最终生成一个Json对象

 输出：
 模块总数： function total number
 最大模块行数：max module execute line number ,    filename
 最小模块行数：min module execute line number ,    filename
 最大圈复杂度数： max matrix complex number,    filename\function name 

 模块信息表：
 文件名            模块数      
 .....                     .......

 模块圈复杂度、扇出数统计表：
 模块名称                                   圈复杂度（超过10）              扇出数（大于7）
 filename/function name             12                                                  9
                ......                                 ......                                              ......


输入： xxx.mts.htm文件
每个testbed生成的最终结果，包含一个xxx.mts.htm文件，存放了软件的质量度量信息
所有的度量信息都是以.c文件为单位，用JSON描述这些信息.
以code.c为例：
{
    1.file name : string ,

    2.reformatted code information for file : (<a id="reformatted code information for file">Reformatted Code Information for File (TM.C)</a>)
    {
        total lines : int
        totalcomments : int
        executeable ref lines : int 
        Non executeable lines: int
        number of procedures : int (模块信息表项)
     },

    3. complexity metrics :  (<a id="complexity metrics">Complexity Metrics (UTIL.C)</a>)
    [
        { Cyclomatic information :  int , function name : string },
        { Cyclomatic information :  int , function name : string },
        { Cyclomatic information :  int , function name : string }
    ],

    4. dataflow information:   (<a id="dataflow information">Dataflow Information (UTIL.C)</a> )
    [
        {       fan out number : int,   function name : string       },
        {       fan out number : int,   function name : string       },
        {       fan out number : int,   function name : string       }
    ]
}

方法：
从每个文件中依次查找
<a id="reformatted code information for file">Reformatted Code Information for File (TM.C)</a>
<a id="complexity metrics">Complexity Metrics (UTIL.C)</a>
<a id="dataflow information">Dataflow Information (UTIL.C)</a>


TODO:
考虑用jison来存储这些信息，并发送给C#生成端程序，这样保证数据库格式与C#分离，在数据库格式修改的情况下，不需要修改C#软件
同理，规则分析的结果也通过jison发给C#生成端
 """


import re
from urllib.request import urlopen
#from    urllib import urlopen
from    bs4 import BeautifulSoup
import json
from config import _config_data



class reformated_code_information(object):
    def __init__(self, totaline=0, totalcomment=0, executeablelines=0, nonexecuteablelines=0, numberOfprocedure=0):
        self.total_line_number = totaline 
        self.total_comments_number = totalcomment
        self.executeable_ref_lines = executeablelines
        self.Non_executeable_lines = nonexecuteablelines 
        self.number_of_procedure = numberOfprocedure
            
    def __str__(self):
        str_output = "{}  {}  {} {} {}".format(self.total_line_number, self.total_comments_number, self.executeable_ref_lines, self.Non_executeable_lines, self.number_of_procedure)
        return str_output



class function_complexity(object):
    def __init__(self, fun_name = '' , cycl_info = 0):
        self.function_name = fun_name
        self.Cyclomatic_information = cycl_info

class function_fanout(object):
    def __init__(self, func_name = '', fout = 0):
        self.function_name = func_name
        self.fanout = fout

class function_procedure_info(object):
    def __init__(self, func_name, executable_lines):
        self.executable_lines = executable_lines
        self.function_name = func_name

class metrix_report(object):
    """
    处理一个源文件的metrix report
    """
    def __init__(self):
        self.filename = ''
        self.reformated_code_information = reformated_code_information()
        self.complexity_metrics = [] #list of function complexity
        self.fanout_info = [] #list of function fanout
        self.func_procedure_info = [] #list of function procedure info
        return 

    def get_max_min_lines(self):
        """
        获取所有本文件最大/最小模块行号
        """
        execute_lines = [e.executable_lines for e in self.func_procedure_info]
        execute_lines.sort(reverse = True)

        return  execute_lines[0], execute_lines[-1]

    def write_to_table(self, source_line_table, metrix_table, complexity_fanout_table):
        """
        将本文件的metrix信息写入表格
        """

        #填写各文件代码行数
        index= len(source_line_table.rows)
        new_cells = source_line_table.add_row().cells
        new_cells[0].text = str(index)
        new_cells[1].text = self.filename
        new_cells[2].text = str(self.reformated_code_information.executeable_ref_lines)
        #填写静态质量度量
        index = len(metrix_table.rows)
        new_cells = metrix_table.add_row().cells
        (max_line, min_line)= self.get_max_min_lines()
        metrix_tuple = (str(index), self.filename, str(self.reformated_code_information.number_of_procedure), '/'.join(str(e)  for e in (max_line, min_line )))
        
        for i, e in enumerate(metrix_tuple):
            new_cells[i].text = e

        #填写复杂度、扇出数表格
        index = len(complexity_fanout_table.rows)
        
        complexity_list = []
        for e in self.complexity_metrics:
            row_list = [index, e.function_name, e.Cyclomatic_information]
            for m in self.fanout_info:
                if m.function_name == e.function_name:
                    row_list.append(m.fanout)
            if row_list[2] > 10 or row_list[3] > 7:
                complexity_list.append(row_list)
            index += 1
            
        for e in complexity_list:
            new_cells = complexity_fanout_table.add_row().cells
            for i, m in enumerate(e):
                new_cells[i].text = str(m)
        return

    def store_db(self, db_obj):
        """将一个软件的testbed 度量分析结果存入DB
        mainly fill two tables:
        source_file_info:
        complexity_metrics_info:
        """
        userid , proid = db_obj.get_userid_projectid()

        #fill source_file_info table
        row_tuple = (proid,  self.filename,  self.reformated_code_information.total_line_number,  self.reformated_code_information.total_comments_number , \
                    self.reformated_code_information.executeable_ref_lines , self.reformated_code_information.number_of_procedure )
        file_id = db_obj.insert_LDRA_metrics(row_tuple)

        #fill complexity_metrics_info table
        for e in self.complexity_metrics:        
            sql_str = '''insert into complextity_metrics_info (file_id,  funtion_name, cyclomatic, fan_out)
                values (?,?,?,?) '''
            sql_tuple = (file_id, e.function_name , e.Cyclomatic_information, '') 
            db_obj.execute_sql_stm(sql_str, sql_tuple)
        
        for e in self.fanout_info:
            sql_str = "update complextity_metrics_info set fan_out = ? where funtion_name = ?"
            sql_tuple = (e.fanout , e.function_name)
            db_obj.execute_sql_stm(sql_str, sql_tuple)

        return

class process_metrix_repot(object):
    def __init__(self):
        self.fileurl = ''
        self.html = None
        self.bsObj = None
        self.match_digital_string = '\d+\.?\d*' #匹配数字
        self.match_file_name = '\([^()]*\)'    #匹配括号内源文件名称
        self.total_info_dict = {} #filename , metrix_report class instance
        self.seri_json = None

    def get_info_from_content_table(self, content_table):
        tr_list = content_table.find_all('tr')
        if len(tr_list) > 4:
            tr_list = tr_list[3:-2] #前面3行是固定的无用内容，后面2行是无用的内容
        else:
            tr_list = tr_list[3:]

        content_list = []
        for tr in tr_list:
            td_list = tr.find_all('td')
            tr_string = [s.string.strip() for s in td_list]
            content_list.append(tr_string)
        return content_list

    def get_reformated_info(self, content_table):
        """
        获取软件基本信息，包括模块数量
        处理<a id="reformatted code information for file">tag下面的表格信息，表格内如如下
        <table bgcolor="#ECE2E2" width="100%">
        <tbody>
        <tr align="LEFT"><th> File </th><th> Total Ref.  </th><th> Total
        </th><th> Executable </th><th> Non-Executable </th><th> Number of
        </th><th> Total </th><th> Expansion </th></tr>
        <tr align="LEFT"><th> &nbsp; </th><th> <u> Lines </u> </th><th> <u>
        Comments </u> </th><th> <u> Ref.  Lines </u> </th><th> <u> Ref.  Lines
        </u> </th><th> <u> Procedures </u> </th><th> <u> Src.  Lines </u>
        </th><th> <u> Factor </u> </th></tr>
        <tr align="LEFT"><th colspan="8"> &nbsp; </th></tr>
        <tr align="LEFT"><td> Total for UTIL.C </td><td><font color="#008000">
        538 (P) </font></td><td><font color="#008000"> 92 (P) (17%)
        </font></td><td><font color="#008000"> 245 (P) (46%)
        </font></td><td><font color="#008000"> 201 (P) (37%)
        </font></td><td><font color="#008000"> 6 (P) </font></td><td><font
        color="#008000"> 315 (P) </font></td><td><font color="#008000"> 1.71
        (P) </font></td></tr>
        </tbody></table>
        """
        content_list = self.get_info_from_content_table(content_table)
        numberlist = [n for n in content_list[0][1 : 6]]
        #使用正则表达式把真实数字从numberlist中提取出来
        real_number = [ int(self.script_regex.findall(numberstring)[0]) for numberstring in numberlist ]
        #一个很方便的列表初始化函数参数列表的方法 function(*list_name)
        reformat_code_info = reformated_code_information(*real_number)
        
        return reformat_code_info

    def get_metrix_complextity_info(self,  content_table):
        """
        获取圈复杂度
        处理<a id="complexity metrics">Complexity Metrics (UTIL.C)</a>下面的信息
        """
        comlextity_list= []
        content_list = self.get_info_from_content_table(content_table)
        for s in content_list:
            function_name , complextity_num = s[0], s[2]
            number_mattched = self.script_regex.findall(complextity_num)
            if len(number_mattched) != 0: #testbed 报告每10个函数有一个空行
                complextity_num = int(number_mattched[0])
                comlextity_list.append((function_name , complextity_num ))
        return comlextity_list

    def get_procedure_info(self, content_table):
        procedure_info_list= []
        content_list = self.get_info_from_content_table(content_table)
        for s in content_list:
            function_name , executeable_lines = s[0], s[1]
            number_mattched = self.script_regex.findall(executeable_lines)
            if len(number_mattched) != 0: #testbed 报告每10个函数有一个空行
                executeable_lines = int(number_mattched[0])
                procedure_info_list.append((function_name , executeable_lines ))
        return procedure_info_list

    def get_data_flow_info(self,  content_table):
        data_flow_list= []
        content_list = self.get_info_from_content_table(content_table)
        for s in content_list:
            function_name , fanout_num = s[0], s[3]
            number_mattched = self.script_regex.findall(fanout_num)
            if len(number_mattched) != 0: #testbed 报告每10个函数有一个空行
                fanout_num = int(number_mattched[0])
                data_flow_list.append((function_name , fanout_num))
        return data_flow_list


    def extract_file_name(self, reformat_tag):
        """
        从tag中获取源文件名称
        """
        filename = reformat_tag.string
        filename =self.filename_regex.findall(filename)[0]
        filename = filename[1:-1]
        return filename

    def trans_to_JSON(self):
            """
            将metrix_report转换为JSON对象
            复杂对象如何转换成JSON :        seri_json = json.dumps(v, default=lambda x : x.__dict__)
            对象格式：
            {
	            "filename": "CGEN.C"
	            "reformated_code_information": 
		            {
			            "total_line_number": 689, 
			            "total_comments_number": 182, 
			            "Non_executeable_lines": 250, 
			            "executeable_ref_lines": 257, 
			            "number_of_procedure": 4
		            }, 
	            "complexity_metrics": 
		            [
			            {"Cyclomatic_information": 12, "function_name": "genStmt"}, 
			            {"Cyclomatic_information": 16, "function_name": "genExp"}, 
			            {"Cyclomatic_information": 4, "function_name": "cGen"}, 
			            {"Cyclomatic_information": 1, "function_name": "codeGen"}
		            ], 
	            "fanout_info": 
		            [
			            {"fanout": 9, "function_name": "genStmt"}, 
			            {"fanout": 5, "function_name": "genExp"}, 
			            {"fanout": 3, "function_name": "cGen"}, 
			            {"fanout": 8, "function_name": "codeGen"}
		            ], 
            }
            """
            if self.seri_json is not None:
                return self.seri_json

            with open("seri_to_json.dat", "w") as writefile:
                for k, v in self.total_info_dict.iteritems():
                    self.seri_json = json.dumps(v, default=lambda x : x.__dict__)
                    #debug
                    writefile.write("%s\n" % seri_json)
                    decode_obj = json.loads(seri_json)
                    x = metrix_report()
                    x.__dict__ = decode_obj
            return

    def store_matrix_to_docx(self, docx_obj):
        """根据软件的testbed 度量分析结果，生成doxc文档
        """
        table_list = document.tables

        #should not happen
        if len(table_list) != 1:
            pass

        #open line number table
        cell = table_list[0].cell(4,0) 
        source_line_table = cell.tables[0]
        
        #open metrics table
        cell = table_list[0].cell(5,0) 
        metrix_table = cell.tables[0]

        #open complexity and fanout table
        cell = table_list[0].cell(6,0) 
        complexity_fanout_table = cell.tables[0]

        #fill docx
        for k, v in self.total_info_dict.items():
            v.write_to_table(source_line_table, metrix_table, complexity_fanout_table)

        docx_obj.save('demo.docx')
        return

    def store_matrix_to_db(self, db_obj):
        """将一个软件的testbed 度量分析结果存入DB
        """
        if _config_data['__debug'] == "true":
            db_obj.clear_table('source_file_info')
            db_obj.clear_table("complextity_metrics_info")

        for k, v in self.total_info_dict.iteritems():
            v.store_db(db_obj)

        db_obj.commit()    
        return         

    def analyse_html(self,file_url):
            """
            脚本文件主函数
            处理一个htlm文件
            将文件中规则违背情况抽取出来，放入列表，并返回
    
            Args:
                file_url: html文件路径
       
            Returns:
                如果这个文件包含规则违背情况，则返回详细情况的list
                如果这个文件不包含规则违背情况，则返回空表
            方法：
                从每个文件中依次查找
                <a id="reformatted code information for file">Reformatted Code Information for File (TM.C)</a>
                <a id="procedure information">Procedure Information (UTIL.C)</a>
                <a id="complexity metrics">Complexity Metrics (UTIL.C)</a>
                <a id="dataflow information">Dataflow Information (UTIL.C)</a>
            """
            self.fileurl = file_url
            self.html = urlopen(file_url)
            self.bsObj = BeautifulSoup(self.html.read(), features="html.parser") 
            self.script_regex = re.compile(self.match_digital_string)
            self.filename_regex = re.compile(self.match_file_name)

            #process all reformatted code information
            #list of reformated_code_information class instance
            reformat_codelist = []
            reformat_title_list = self.bsObj.find_all(u'a', id = u'reformatted code information for file')
            with open("temp_reformatted_code_info.dat",'w') as write_file:
                for reformat_tag in reformat_title_list:
                    #从tag的context获取filename
                    filename = self.extract_file_name(reformat_tag)

                    #add to total info dict for future useage
                    one_file_info =   metrix_report()
                    one_file_info.filename = filename
                    self.total_info_dict[filename] = one_file_info

                    write_file.write("%s\n" % reformat_tag)
                    for e in reformat_tag.parent.next_siblings:
                        if e.name == 'center':
                            content_table = e.find_all('table')[1] #获取第二个table
                            #获取第二个table中的tr
                            reformat_code_info = self.get_reformated_info(content_table)
                            #store
                            one_file_info.reformated_code_information = reformat_code_info
                            #DEBUG
                            reformat_codelist.append(reformat_code_info)
                            write_file.write("%s\n" % reformat_code_info)
                            write_file.flush()
                            break

            #process procedure information
            procedure_info_list = []
            procedure_title_list = self.bsObj.find_all(u'a', id = u'procedure information')
            with open('temp_procedure_info.dat', 'w') as write_file:
                for procedure_tag in procedure_title_list:
                    #从tag的context获取filename
                    filename = self.extract_file_name(procedure_tag)
                    if filename in self.total_info_dict:
                        one_file_info = self.total_info_dict[filename]

                    write_file.write("%s\n" % procedure_tag)
                    for e in procedure_tag.parent.next_siblings:
                        if e.name == 'center':
                            content_table = e.find_all('table')[1] #获取第二个table
                            #获取第二个table中的tr
                            procedure_info = self.get_procedure_info(content_table)
                            for e in procedure_info:
                                procedure = function_procedure_info(*e)
                                one_file_info.func_procedure_info.append(procedure)
                            
                            #debug
                            procedure_info_list.append(procedure_info)
                            write_file.write("%s\n" % procedure_info)
                            #write_file.write('%s\n' % content_table)
                            write_file.flush()
                            break

            #dataflow information
        #    3. dataflow information:   (<a id="dataflow information">Dataflow Information (UTIL.C)</a> )
        #    [
        #        {       function name : string,      fan out number : int    },
        #        {       function name : string,      fan out number : int    },
        #        {       function name : string,      fan out number : int    },
        #    ]

            dataflow_list= []
            data_flow_title_list = self.bsObj.find_all(u'a', id = u'dataflow information')
            with open("data_flow_info.dat",'w') as write_file:
                for data_flow_tag in data_flow_title_list:
                    #从tag的context获取filename
                    filename = self.extract_file_name(data_flow_tag)
                    if filename in self.total_info_dict :
                        one_file_info = self.total_info_dict[filename]

                    write_file.write("%s\n" % data_flow_tag)
                    for e in data_flow_tag.parent.next_siblings:
                        if e.name == 'center':
                            content_table = e.find_all('table')[1] #获取第二个table
                            #获取第二个table中的tr
                            data_flow_info = self.get_data_flow_info(content_table)
                            for e in data_flow_info:
                                data_flow = function_fanout(*e)
                                one_file_info.fanout_info.append(data_flow)
                            
                            #debug
                            dataflow_list.append(data_flow_info)
                            write_file.write("%s\n" % data_flow_info)
                            #write_file.write('%s\n' % content_table)
                            write_file.flush()
                            break
            pass

            #process complexity metrics                    
            complexity_metrics_list = []
            complextity_title_list = self.bsObj.find_all(u'a', id = u'complexity metrics')
            with open("complexity_metrics_info.dat",'w') as write_file:
               for complextity_tag in complextity_title_list:
                    filename = self.extract_file_name(complextity_tag)
                    if filename in self.total_info_dict :
                        one_file_info = self.total_info_dict[filename]
                    
                    write_file.write("%s\n" % complextity_tag)
                    for e in complextity_tag.parent.next_siblings:
                        if e.name == 'center':
                            content_table = e.find_all('table')[1] #获取第二个table
                            metrix_info = self.get_metrix_complextity_info(content_table)
                            for e in metrix_info:
                                l = function_complexity(*e)
                                one_file_info.complexity_metrics.append(l)
                            
                            #debug
                            complexity_metrics_list.append(metrix_info)
                            write_file.write("%s\n" % '  '.join(str(e) for e in metrix_info))
                            write_file.flush()
                            break
            return
                    
from store_db_sqlit3 import process_db
from config import _config_data
from docx import Document
from docx.shared import Inches

if __name__ == '__main__':
    
    dev_location = _config_data['dev_location']

    if dev_location == 'home':
        html = u"file:///C:/Users/Administrator/Documents/code/project_from_github/analyse_html/example_tbwrkfls/example.mts.htm"
    else:
        html = u"file:///C:/LDRA_Workarea/example_tbwrkfls/example.mts.htm"

    report = process_metrix_repot()
    report.analyse_html(html)
    
    #以模板为基础，生成度量结果文档
    document = Document(u'质量度量.docx')
    report.store_matrix_to_docx(document)
    #produce doc
    
    #report.trans_to_JSON()
    #db_obj = process_db()
    #report.store_matrix_to_db(db_obj)
